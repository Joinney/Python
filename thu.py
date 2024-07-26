from tkinter import *  # Nhập tất cả các đối tượng và hàm từ thư viện Tkinter
from tkinter import messagebox  # Nhập hàm messagebox từ thư viện Tkinter, để hiển thị các hộp thoại thông báo
from PIL import Image, ImageTk  # Nhập các modul Image và ImageTk từ thư viện PIL, để làm việc với hình ảnh trong Tkinter
import random  # Nhập thư viện random, để sử dụng các chức năng liên quan đến số ngẫu nhiên
import smtplib  # Nhập thư viện smtplib, để gửi email từ Python thông qua giao thức SMTP
from email.mime.multipart import MIMEMultipart  # Nhập lớp MIMEMultipart từ thư viện email.mime.multipart, để tạo và gửi email với các phần tử đa phần (multipart)
from email.mime.text import MIMEText  # Nhập lớp MIMEText từ thư viện email.mime.text, để tạo và gửi email với các phần tử văn bản
import numpy as np  # Nhập thư viện numpy, để làm việc với các mảng và ma trận, cũng như các hàm toán học có hiệu suất cao
from tkinter import Tk, Label  # Nhập các lớp Tk, Label và Canvas từ thư viện Tkinter
import matplotlib.pyplot as plt  # Nhập thư viện matplotlib.pyplot, để vẽ biểu đồ và hình ảnh trong Python
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg  # Nhập lớp FigureCanvasTkAgg từ thư viện matplotlib.backends.backend_tkagg, để chuyển đổi biểu đồ thành một widget Tkinter
from PIL import Image, ImageTk, ImageDraw
import threading,time
from datetime import datetime
import tkinter as tk
from tkinter import ttk, messagebox, GROOVE
import pyodbc
from fpdf import FPDF
from collections import defaultdict
from collections import defaultdict
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

#========================================= LOADING ỨNG DỤNG =========================================
class LoadingScreen:
    def __init__(self, master):
        # Khởi tạo loading screen với các thành phần chính
        self.master = master
        self.master.title("Loading!")  # Đặt tiêu đề cửa sổ
        self.master.geometry("500x300")  # Đặt kích thước cửa sổ

        # Load background image
        # Tải hình nền và chuyển đổi thành đối tượng PhotoImage
        self.background_image = Image.open("loading.png")
        self.background_photo = ImageTk.PhotoImage(self.background_image)
        
        # Tạo Canvas để đặt hình nền
        self.canvas = tk.Canvas(self.master, width=500, height=300)
        self.canvas.pack(fill="both", expand=True)

        # Đặt hình nền lên Canvas
        self.canvas.create_image(0, 0, image=self.background_photo, anchor="nw")

        # Tạo label chứa tiêu đề chào mừng
        self.label_title = tk.Label(self.master, text="WELCOME! BÁCH HÓA MINI", bg="green", fg="yellow", font=("Arial", 20))
        self.label_title.place(relx=0.5, rely=0.1, anchor=tk.CENTER)
        self.vibrate_id = self.master.after(100, self.vibrate_title)  # Kích hoạt hiệu ứng rung cho tiêu đề

        # Tạo thanh tiến trình
        self.progress_bar = ttk.Progressbar(self.master, orient="horizontal", length=400, mode="determinate")
        self.progress_bar.place(relx=0.5, rely=0.9, anchor=tk.CENTER)  # Di chuyển thanh tiến trình xuống dưới
        self.progress_bar_label = tk.Label(self.master, text="Loading 0%", bg="green", fg="yellow", font=("Arial", 12))
        self.progress_bar_label.place(relx=0.5, rely=0.95, anchor=tk.CENTER)  # Đặt label dưới thanh tiến trình
        self.progress_id = self.master.after(50, self.animate_progress)  # Bắt đầu hiệu ứng tiến trình

    def vibrate_title(self):
        # Tạo hiệu ứng rung cho tiêu đề
        x_offset = random.uniform(-1, 1)
        y_offset = random.uniform(-1, 1)
        self.label_title.place(relx=0.5 + x_offset/100, rely=0.1 + y_offset/100, anchor=tk.CENTER)
        self.vibrate_id = self.master.after(100, self.vibrate_title)  # Lặp lại sau 100ms

    def animate_progress(self):
        # Tạo hiệu ứng tiến trình
        value = self.progress_bar["value"]
        if value < 100:
            self.progress_bar["value"] += 1
            self.progress_bar_label.config(text=f"Loading {int(value)}%")
            self.progress_id = self.master.after(50, self.animate_progress)  # Lặp lại sau 50ms
        else:
            self.on_progress_complete()

    def on_progress_complete(self):
        # Hủy tất cả các after callbacks
        self.master.after_cancel(self.vibrate_id)
        self.master.after_cancel(self.progress_id)
        # Chuyển sang trang mới khi tiến trình hoàn tất
        self.master.destroy()  # Đóng cửa sổ hiện tại
        new_root = tk.Tk()
        new_app = Login_App(new_root)
        new_root.mainloop()






#========================================= Cửa sổ đăng nhập =========================================
class Login_App:
    def __init__(self, root):
        self.root = root
        self.root.geometry("700x400")
        self.root.maxsize(width=700, height=400)
        self.root.minsize(width=700, height=400)
        self.root.title("Đăng nhập")
        
        # Load background image
        self.bg_image = Image.open("shop.png")
        self.bg_photo = ImageTk.PhotoImage(self.bg_image)
        self.bg_label = Label(self.root, image=self.bg_photo)
        self.bg_label.place(x=0, y=0, relwidth=1, relheight=1)

        self.setup_login()

    def setup_login(self):
        bg_color = "#f2f2f2"
        fg_color = "black"
        label_font = ("Arial", 12)
        entry_font = ("Arial", 12)
        button_font = ("Arial", 12, "bold")

        self.username = StringVar()
        self.password = StringVar()
        
        label_title = Label(self.root, text=" WELLCOM! BÁCH HÓA MINI", bg="green", fg="yellow", font=("Arial", 20))
        label_title.place(relx=0.5, rely=0.1, anchor=CENTER)

        img1 = Image.open("logo.png")
        img1 = img1.resize((50, 50), Image.ADAPTIVE)

        mask = Image.new("L", (50, 50), 0)
        draw = ImageDraw.Draw(mask)
        draw.ellipse((0, 0, 50, 50), fill=255)

        circle_img = Image.new("RGBA", (50, 50), (0, 128, 0, 255))
        circle_img.paste(img1, (0, 0), mask=mask)

        self.photoimg1 = ImageTk.PhotoImage(circle_img)

        b1 = Button(self.root, image=self.photoimg1, borderwidth=0, bg="green")
        b1.place(x=100, y=15)

        self.login_frame = LabelFrame(self.root, font=("time new roman", 12, "bold"), fg="gold", bg="white", relief=GROOVE, bd=10)
        self.login_frame.place(relx=0.5, rely=0.5, anchor=CENTER)

        Label(self.login_frame, text="Tên ID:", bg=bg_color, fg=fg_color, font=label_font).grid(row=0, column=0, padx=10, pady=5, sticky="e")
        Entry(self.login_frame, textvariable=self.username, font=entry_font, bd=2, relief=GROOVE).grid(row=0, column=1, padx=10, pady=5)
        Label(self.login_frame, text="Mật khẩu:", bg=bg_color, fg=fg_color, font=label_font).grid(row=1, column=0, padx=10, pady=5, sticky="e")
        Entry(self.login_frame, textvariable=self.password, font=entry_font, bd=2, relief=GROOVE, show="*").grid(row=1, column=1, padx=10, pady=5)
        
        btn_login = Button(self.login_frame, text="Đăng nhập", width=10, command=self.login, font=button_font, bg="#4CAF50", fg="white", relief=RAISED)
        btn_login.grid(row=2, columnspan=2, pady=10)

        

        btn_register = Button(self.login_frame, text="Đăng ký", width=10, command=self.register, font=button_font, bg="#4CAF50", fg="white", relief=RAISED)
        btn_register.grid(row=3, columnspan=2, pady=10)

    def login(self):
        username = self.username.get()
        password = self.password.get()
        try:
            with open("DANGNHAP.txt", "r") as file:
                for line in file:
                    user, pwd = line.strip().split(",")
                    if username == user and password == pwd:
                        messagebox.showinfo("Đăng nhập", "Đăng nhập thành công.")
                        self.root.destroy()
                        app = MainApp(Tk())
                        app.root.mainloop()
                        return
                messagebox.showerror("Lỗi Đăng nhập", "Tên người dùng hoặc mật khẩu không đúng.")
        except FileNotFoundError:
            messagebox.showerror("Lỗi Đăng nhập", "Tệp người dùng không tồn tại. Vui lòng đăng ký trước.")
    def open_forgot_password_window(self, event):
        # Create a new window for forgot password
        forgot_password_window = Toplevel(self.root)
        forgot_password_window.title("Quên mật khẩu")

        # Label and Entry for email address
        lbl_email = Label(forgot_password_window, text="Nhập địa chỉ email của bạn:", font=("Arial", 12))
        lbl_email.pack(pady=10)
        entry_email = Entry(forgot_password_window, font=("Arial", 12))
        entry_email.pack(pady=5)

        # Button to send password reset link
        btn_send_reset_link = Button(forgot_password_window, text="Gửi mã đặt lại", command=lambda: self.send_reset_link(entry_email.get()), font=("Arial", 12))
        btn_send_reset_link.pack(pady=10)

    def send_reset_link(self, email):
        # Implement logic to send reset link to the provided email address
        # For demonstration, I'll just show a message box
        messagebox.showinfo("Quên mật khẩu", f"Một mã đặt lại mật khẩu đã được gửi đến '{email}'.")
    
    
    def register(self):
        self.root.withdraw()
        register_window = DangKy(self.root)
        self.root.wait_window(register_window.register_window)
        self.root.deiconify()


#============================================================ Cửa sổ đăng ký =====================================================
class DangKy:
    def __init__(self, parent):
        self.parent = parent
        self.register_window = Toplevel(parent)
        self.register_window.title("Đăng ký")
        self.register_window.geometry("700x400")
        # Load background image
        bg_image = Image.open("shop.png")
        self.bg_photo = ImageTk.PhotoImage(bg_image)
        self.bg_label = Label(self.register_window, image=self.bg_photo)
        self.bg_label.place(x=0, y=0, relwidth=1, relheight=1)
        
        self.setup_register()

        # Nhãn tiêu đề chính
        label_title2 = Label(
            self.register_window,
            text=" WELLCOM! BÁCH HÓA MINI",
            bg="green",
            fg="yellow",
            font=("Arial", 20)
        )
        label_title2.place(relx=0.5, rely=0.1, anchor=CENTER)

        # Tải và thay đổi kích thước hình ảnh
        img1 = Image.open("logo.png")
        img1 = img1.resize((50, 50), Image.ADAPTIVE)

        # Tạo hình tròn trong suốt
        mask = Image.new("L", (50, 50), 0)
        draw = ImageDraw.Draw(mask)
        draw.ellipse((0, 0, 50, 50), fill=255)

        # Tạo hình ảnh mới có kích thước và định dạng tương tự như hình ảnh gốc
        circle_img = Image.new("RGBA", (50, 50), (0, 128, 0, 255))

        # Áp dụng hình ảnh gốc lên hình tròn trong suốt
        circle_img.paste(img1, (0, 0), mask=mask)

        # Tạo PhotoImage từ hình ảnh vừa tạo
        self.photoimg1 = ImageTk.PhotoImage(circle_img)

        # Nút với hình ảnh hình tròn và màu nền xanh
        b1 = Button(self.register_window, image=self.photoimg1, borderwidth=0, bg="green")
        b1.place(x=100, y=15)

    def setup_register(self):
        bg_color = "#f2f2f2"
        fg_color = "black"
        label_font = ("Arial", 12)
        entry_font = ("Arial", 12)
        button_font = ("Arial", 12, "bold")

        # Biến để lưu trữ tên người dùng, mật khẩu, email và OTP
        self.username = StringVar()
        self.password = StringVar()
        self.email = StringVar()
        self.otp = StringVar()

        # Khung cho các phần đăng ký
        register_frame = LabelFrame(self.register_window, font=("time new roman", 12, "bold"), fg="gold", bg="white", relief=GROOVE, bd=10)
        register_frame.place(relx=0.5, rely=0.5, anchor=CENTER)

        # Nhãn và trường nhập cho tên người dùng, mật khẩu và email
        Label(
            register_frame,
            text="Tên ID:",
            bg=bg_color,
            fg=fg_color,
            font=label_font
        ).grid(row=0, column=0, padx=10, pady=5, sticky="e")

        Entry(
            register_frame,
            textvariable=self.username,
            font=entry_font,
            bd=2,
            relief=GROOVE
        ).grid(row=0, column=1, padx=10, pady=5)

        Label(
            register_frame,
            text="Mật khẩu:",
            bg=bg_color,
            fg=fg_color,
            font=label_font
        ).grid(row=1, column=0, padx=10, pady=5, sticky="e")

        Entry(
            register_frame,
            textvariable=self.password,
            font=entry_font,
            bd=2,
            relief=GROOVE,
            show="*"
        ).grid(row=1, column=1, padx=10, pady=5)

        Label(
            register_frame,
            text="Email:",
            bg=bg_color,
            fg=fg_color,
            font=label_font
        ).grid(row=2, column=0, padx=10, pady=5, sticky="e")

        Entry(
            register_frame,
            textvariable=self.email,
            font=entry_font,
            bd=2,
            relief=GROOVE
        ).grid(row=2, column=1, padx=10, pady=5)

        # Nút để gửi mã OTP
        self.btn_send_otp = Button(
            register_frame,
            text="Gửi mã OTP",
            width=15,
            command=self.send_otp,
            font=button_font,
            bg="#4CAF50",
            fg="white",
            relief=RAISED
        )
        self.btn_send_otp.grid(row=3, columnspan=2, pady=10)

        # Nhãn và trường nhập cho mã OTP
        Label(
            register_frame,
            text="Nhập mã OTP:",
            bg=bg_color,
            fg=fg_color,
            font=label_font
        ).grid(row=4, column=0, padx=10, pady=5, sticky="e")

        Entry(
            register_frame,
            textvariable=self.otp,
            font=entry_font,
            bd=2,
            relief=GROOVE
        ).grid(row=4, column=1, padx=10, pady=5)

        # Nút để đăng ký
        btn_register = Button(
            register_frame,
            text="Đăng ký",
            width=10,
            command=self.register_user,
            font=button_font,
            bg="#4CAF50",
            fg="white",
            relief=RAISED
        )
        btn_register.grid(row=5, columnspan=2, pady=10)

    def send_otp(self):
        self.generated_otp = self.generate_otp()  # Tạo mã OTP mới

        sender_email = "zBachHoaminiz@gmail.com"
        sender_password = "mnzw yarr lynw ubma"  # Mật khẩu ứng dụng Gmail
        receiver_email = self.email.get()

        msg = MIMEMultipart()
        msg['From'] = sender_email
        msg['To'] = receiver_email
        msg['Subject'] = "BÁCH HÓA MINI XIN CHÀO BẠN!"

        body = f"Mã OTP Đăng Ký Của Bạn Là: {self.generated_otp}"
        msg.attach(MIMEText(body, 'plain'))

        try:
            # Kết nối và gửi email
            server = smtplib.SMTP('smtp.gmail.com', 587)
            server.starttls()  # Kích hoạt TLS
            server.login(sender_email, sender_password)
            server.sendmail(sender_email, receiver_email, msg.as_string())
            server.quit()

            messagebox.showinfo("Thông báo", "Mã OTP đã được gửi đến email của bạn.")
            # Khóa nút gửi và bắt đầu đếm ngược trước khi cho phép gửi lại OTP
            countdown_thread = threading.Thread(target=self.start_countdown)
            countdown_thread.start()
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể gửi email: {e}")
            self.btn_send_otp.config(state="normal")

    def start_countdown(self):
        # Đếm ngược 30 giây
        for remaining in range(30, -1, -1):
            if self.register_window.winfo_exists():  # Kiểm tra nếu cửa sổ vẫn tồn tại
                self.btn_send_otp.config(text=f"Gửi mã OTP ({remaining}s)", state="disabled")
                time.sleep(1)
            else:
                return

        if self.register_window.winfo_exists():
            self.btn_send_otp.config(text="Gửi mã OTP", state="normal")

    def generate_otp(self):
        return str(random.randint(100000, 999999))  # Trả về số ngẫu nhiên 6 chữ số

    def register_user(self):
        username = self.username.get()
        password = self.password.get()
        otp_input = self.otp.get()

        if username and password and otp_input:  # Kiểm tra nếu các trường đã được điền
            if otp_input == self.generated_otp:  # Kiểm tra nếu OTP khớp
                try:
                    with open("DANGNHAP.txt", "a") as file:
                        file.write(f"{username},{password}\n")
                    messagebox.showinfo("Đăng ký", "Đăng ký thành công.")
                    self.register_window.destroy()  # Đóng cửa sổ sau khi đăng ký thành công
                except Exception as e:
                    messagebox.showerror("Lỗi Đăng ký", f"Không thể đăng ký: {e}")
            else:
                messagebox.showerror("Lỗi Đăng ký", "Mã OTP không khớp. Vui lòng thử lại.")  # Thông báo khi OTP không khớp
        else:
            messagebox.showerror("Lỗi Đăng ký", "Vui lòng nhập đầy đủ thông tin và mã OTP.")  # Khi các trường không đầy đủ


#==============================cửa sổ danh thu======================================================================cửa sổ tổng danh thu======================================================================================#            
class DanhThu:
    def __init__(self, root):
        self.root = root
        self.root.geometry("1100x700")
        self.root.title("Tổng doanh thu")
        self.current_view = 'items'  # Track the current view
        self.root.configure(background="green")
        self.setup_page()

    def setup_page(self):
        self.login_frame = LabelFrame(self.root, font=("time new roman", 12, "bold"), fg="gold", bg="white", relief="groove", bd=10)
        self.login_frame.place(relx=0.5, rely=0.5, anchor=CENTER)
        
        self.btn_toggle_chart = Button(self.root, text="Hiển thị biểu đồ doanh thu theo tháng", bg='green', fg='white', bd=5, relief = GROOVE, command=self.toggle_chart)
        self.btn_toggle_chart.pack(pady=10)

        self.display_item_sales_chart()

    def display_item_sales_chart(self):
        try:
            conn = pyodbc.connect(
                Driver='{SQL Server}',
                SERVER='TEO-PC\\SQLEXPRESS',  # Replace with your server name
                Database='QLBH',  # Replace with your database name
                Trusted_Connection='yes'
            )

            cursor = conn.cursor()
            cursor.execute("SELECT Bread, Candy, Hamburger, Hotdog, Sandwich, Wheat, FoodOil, Salt, Rice, Sugar, Gatorade, Juice, Coke, Waffer, Biscuits FROM DANHSACH")
            rows = cursor.fetchall()

            items = ['Bread', 'Candy', 'Hamburger', 'Hotdog', 'Sandwich', 'Wheat', 'FoodOil', 'Salt', 'Rice', 'Sugar', 'Gatorade', 'Juice', 'Coke', 'Waffer', 'Biscuits']
            sales_data = [sum(row) for row in zip(*rows)]

            fig, ax = plt.subplots(figsize=(10, 6))
            colors = ['skyblue', 'orange', 'green', 'red', 'purple', 'brown', 'pink', 'gray', 'olive', 'cyan', 'magenta', 'yellow', 'blue', 'indigo', 'teal']
            ax.bar(items, sales_data, color=colors)
            ax.set_xlabel('Mặt hàng')
            ax.set_ylabel('Số lượng bán được')
            ax.set_title('Biểu đồ số lượng bán được của mỗi mặt hàng')
            plt.xticks(rotation=45, ha='right')

            for i, sales in enumerate(sales_data):
                ax.text(i, sales, str(sales), ha='center', va='bottom')

            self.clear_canvas()
            self.canvas = FigureCanvasTkAgg(fig, master=self.login_frame)
            self.canvas.draw()
            self.canvas.get_tk_widget().pack()

            conn.close()

        except pyodbc.Error as e:
            print("Lỗi khi kết nối cơ sở dữ liệu:", e)

    def display_monthly_revenue_chart(self):
        try:
            conn = pyodbc.connect(
                Driver='{SQL Server}',
                SERVER='TEO-PC\\SQLEXPRESS',  # Replace with your server name
                Database='QLBH',  # Replace with your database name
                Trusted_Connection='yes'
            )

            cursor = conn.cursor()
            cursor.execute("SELECT Date, Total FROM DANHSACH")
            rows = cursor.fetchall()

            monthly_sales = defaultdict(float)
            for row in rows:
                date, total = row
                month = datetime.strptime(date, '%Y-%m-%d').strftime('%m')
                monthly_sales[month] += float(total)

            sorted_monthly_sales = sorted(monthly_sales.items(), key=lambda x: x[0])
            months, sales = zip(*sorted_monthly_sales)

            fig, ax = plt.subplots(figsize=(10, 6))
            ax.bar(months, sales, color='skyblue')
            ax.set_xlabel('Tháng')
            ax.set_ylabel('Tổng doanh thu ($)')
            ax.set_title('Biểu đồ tổng doanh thu theo tháng')
            plt.xticks(rotation=0, ha='right')

            for i, sale in enumerate(sales):
                ax.text(i, sale, str(sale), ha='center', va='bottom', fontdict={'fontsize': 10})

            self.clear_canvas()
            self.canvas = FigureCanvasTkAgg(fig, master=self.login_frame)
            self.canvas.draw()
            self.canvas.get_tk_widget().pack()

            conn.close()

        except pyodbc.Error as e:
            print("Lỗi khi kết nối cơ sở dữ liệu:", e)

    def toggle_chart(self):
        if self.current_view == 'items':
            self.display_monthly_revenue_chart()
            self.btn_toggle_chart.config(text="Hiển thị biểu đồ số lượng bán được của mỗi mặt hàng", bg='green', fg='white', bd=5, relief = GROOVE)
            self.current_view = 'monthly'
        else:
            self.display_item_sales_chart()
            self.btn_toggle_chart.config(text="Hiển thị biểu đồ doanh thu theo tháng", bg='green', fg='white', bd=5, relief = GROOVE)
            self.current_view = 'items'

    def clear_canvas(self):
        for widget in self.login_frame.winfo_children():
            widget.destroy()
# ==========================================================================================các phần chính của main=================================================================================#
class MainApp: 
    def __init__(self,root):
        self.root = root
        self.root.geometry("1300x700")
        self.root.maxsize(width = 1280,height = 700)
        self.root.minsize(width = 1280,height = 700)
        self.root.title("BÁCH HÓA XANH")

        
        #====================Variables========================#
        self.cus_name = StringVar()
        self.c_phone = StringVar()
        #For Generating Random Bill Numbers
        x = random.randint(1000,9999)
        self.c_bill_no = StringVar()
        #Seting Value to variable
        self.c_bill_no.set(str(x))
        self.e_email=StringVar()
        self.bread = IntVar()
        self.candy = IntVar()
        self.hamburger = IntVar()
        self.hotdog = IntVar()
        self.sandwich = IntVar()
        self.rice = IntVar()
        self.salt = IntVar()
        self.food_oil = IntVar()
        self.wheat = IntVar()
        self.sugar = IntVar()
        self.gatorade = IntVar()
        self.coke = IntVar()
        self.juice = IntVar()
        self.waffer = IntVar()
        self.biscuits = IntVar()
        self.total_food = StringVar()
        self.total_grocery = StringVar()
        self.total_other = StringVar()
        self.tax_cos = StringVar()
        self.tax_groc = StringVar()
        self.tax_other = StringVar()
  

        #===================================
        bg_color = "green"
        fg_color = "white"
        lbl_color = 'white'
        #=================tiêu đề cửa hàng bách hóa xanh=============#
        #Title of App
        title = Label(self.root,text = "BÁCH HÓA MINI",bd = 12,relief = GROOVE,fg = "yellow",bg = bg_color,font=("times new roman",30,"bold"),pady = 3).pack(fill = X)

        #==========================================================================================================================================
        # Tải và thay đổi kích thước hình ảnh
        img1 = Image.open("logo.png")
        img1 = img1.resize((50, 50), Image.ADAPTIVE)

        # Tạo hình tròn trong suốt
        mask = Image.new("L", (50, 50), 0)
        draw = ImageDraw.Draw(mask)
        draw.ellipse((0, 0, 50, 50), fill=255)

        # Tạo hình ảnh mới có kích thước và định dạng tương tự như hình ảnh gốc
        circle_img = Image.new("RGBA", (50, 50), (0, 128, 0, 255))  # Đặt màu xanh cho nền

        # Áp dụng hình ảnh gốc lên hình tròn trong suốt
        circle_img.paste(img1, (0, 0), mask=mask)

        # Tạo PhotoImage từ hình ảnh vừa tạo
        self.photoimg1 = ImageTk.PhotoImage(circle_img)

        # Nút với hình ảnh hình tròn và màu nền xanh
        b1 = Button(self.root, image=self.photoimg1, borderwidth=0, bg="green")
        b1.place(x=400, y=11)
        #======================================================================================================================================

        #==========Customers Frame==========#
        F1 = LabelFrame(text = "Customer Details",font = ("time new roman",12,"bold"),fg = "gold",bg = bg_color,relief = GROOVE,bd = 10)
        F1.place(x = 0,y = 70,relwidth = 1)
        
        #=================Date and Time Display=====================#
        self.date_time_label = Label(self.root, text="", font=("Arial", 14), fg="white", bg="green")
        self.date_time_label.place(x=1050, y=35)
        # Update date and time initially and then every second
        self.update_datetime()
        self.root.after(1000, self.update_datetime)
        
        #===============Customer Name===========#
        cname_lbl = Label(F1,text="Customer Name",bg = bg_color,fg = fg_color,font=("times new roman",15,"bold"))
        cname_lbl.grid(row = 0,column = 0,padx = 10,pady = 5)
        cname_en = Entry(F1,bd = 8,relief = GROOVE,textvariable = self.cus_name)
        cname_en.grid(row = 0,column = 1,ipady = 4,ipadx = 30,pady = 5)

        #=================Customer Phone==============#
        cphon_lbl = Label(F1,text = "Phone No",bg = bg_color,fg = fg_color,font = ("times new roman",15,"bold"))
        cphon_lbl.grid(row = 0,column = 2,padx = 20)
        cphon_en = Entry(F1,bd = 8,relief = GROOVE,textvariable = self.c_phone)
        cphon_en.grid(row = 0,column = 3,ipady = 4,ipadx = 30,pady = 5)

         #===============TEXT EMAIL====================
        email_lbl = Label(F1,text = "Email",bg = bg_color,fg = fg_color,font = ("times new roman",15,"bold"))
        email_lbl.grid(row = 0,column = 4,padx = 20)
        email_en=Entry(F1,bd = 8,relief = GROOVE,textvariable = self.e_email)
        email_en.grid(row = 0,column = 5,ipady = 4,ipadx = 0,pady = 5)


        #====================Customer Bill No==================#
        cbill_lbl = Label(F1,text = "Bill No.",bg = bg_color,fg = fg_color,font = ("times new roman",15,"bold"))
        cbill_lbl.grid(row = 0,column = 6,padx = 20)
        cbill_en = Entry(F1,bd = 8,relief = GROOVE,textvariable = self.c_bill_no)
        cbill_en.grid(row = 0,column = 7,ipadx = 0,ipady = 4,pady = 0)

        save_button = Button(F1,text = "SAVE", font=('Arial', 12, 'bold'), bg='green', fg='white',bd=7,relief = GROOVE, command=self.save_to_database)
        save_button.grid(row =0,column = 9,ipadx = 4,ipady = 4,pady = 5)

        #==================Food Frame=====================#
        F2 = LabelFrame(self.root,text = 'Food',bd = 10,relief = GROOVE,bg = bg_color,fg = "gold",font = ("times new roman",13,"bold"))
        F2.place(x = 0,y = 170,width = 325,height = 380)

        #===========Frame Content======================
        bath_lbl = Label(F2,font = ("times new roman",15,"bold"),fg = lbl_color,bg = bg_color,text = "Bread")
        bath_lbl.grid(row = 0,column = 0,padx = 10,pady = 20)
        bath_en = Entry(F2,bd = 8,relief = GROOVE,textvariable = self.bread)
        bath_en.grid(row = 0,column = 1,ipady = 5,ipadx = 5)

        #================Candy=============================
        face_lbl = Label(F2,font = ("times new roman",15,"bold"),fg = lbl_color,bg = bg_color,text = "Candy")
        face_lbl.grid(row = 1,column = 0,padx = 10,pady = 20)
        face_en = Entry(F2,bd = 8,relief = GROOVE,textvariable = self.candy)
        face_en.grid(row = 1,column = 1,ipady = 5,ipadx = 5)

        #========Hamburger================================
        wash_lbl = Label(F2,font = ("times new roman",15,"bold"),fg = lbl_color,bg = bg_color,text = "Hamburger")
        wash_lbl.grid(row = 2,column = 0,padx = 10,pady = 20)
        wash_en = Entry(F2,bd = 8,relief = GROOVE,textvariable = self.hamburger)
        wash_en.grid(row = 2,column = 1,ipady = 5,ipadx = 5)

        #========Hotdog=============================
        hair_lbl = Label(F2,font = ("times new roman",15,"bold"),fg = lbl_color,bg = bg_color,text = "Hotdog")
        hair_lbl.grid(row = 3,column = 0,padx = 10,pady = 20)
        hair_en = Entry(F2,bd = 8,relief = GROOVE,textvariable = self.hotdog)
        hair_en.grid(row = 3,column = 1,ipady = 5,ipadx = 5)

        #============Sandwich
        lot_lbl = Label(F2,font = ("times new roman",15,"bold"),fg = lbl_color,bg = bg_color,text = "Sandwich")
        lot_lbl.grid(row = 4,column = 0,padx = 10,pady = 20)
        lot_en = Entry(F2,bd = 8,relief = GROOVE,textvariable = self.sandwich)
        lot_en.grid(row = 4,column = 1,ipady = 5,ipadx = 5)

        #==================Grocery Frame=====================#
        F2 = LabelFrame(self.root,text = 'Grocery',bd = 10,relief = GROOVE,bg = bg_color,fg = "gold",font = ("times new roman",13,"bold"))
        F2.place(x = 325,y = 170,width = 325,height = 380)

        #===========Frame Content
        rice_lbl = Label(F2,font = ("times new roman",15,"bold"),fg = lbl_color,bg = bg_color,text = "Rice")
        rice_lbl.grid(row = 0,column = 0,padx = 10,pady = 20)
        rice_en = Entry(F2,bd = 8,relief = GROOVE,textvariable = self.rice)
        rice_en.grid(row = 0,column = 1,ipady = 5,ipadx = 5)

        #=======
        oil_lbl = Label(F2,font = ("times new roman",15,"bold"),fg = lbl_color,bg = bg_color,text = "Food Oil")
        oil_lbl.grid(row = 1,column = 0,padx = 10,pady = 20)
        oil_en = Entry(F2,bd = 8,relief = GROOVE,textvariable = self.food_oil)
        oil_en.grid(row = 1,column = 1,ipady = 5,ipadx = 5)

        #=======
        daal_lbl = Label(F2,font = ("times new roman",15,"bold"),fg = lbl_color,bg = bg_color,text = "Salt")
        daal_lbl.grid(row = 2,column = 0,padx = 10,pady = 20)
        daal_en = Entry(F2,bd = 8,relief = GROOVE,textvariable = self.salt)
        daal_en.grid(row = 2,column = 1,ipady = 5,ipadx = 5)

        #========
        wheat_lbl = Label(F2,font = ("times new roman",15,"bold"),fg = lbl_color,bg = bg_color,text = "Wheat")
        wheat_lbl.grid(row = 3,column = 0,padx = 10,pady = 20)
        wheat_en = Entry(F2,bd = 8,relief = GROOVE,textvariable = self.wheat)
        wheat_en.grid(row = 3,column = 1,ipady = 5,ipadx = 5)

        #============
        sugar_lbl = Label(F2,font = ("times new roman",15,"bold"),fg = lbl_color,bg = bg_color,text = "Sugar")
        sugar_lbl.grid(row = 4,column = 0,padx = 10,pady = 20)
        sugar_en = Entry(F2,bd = 8,relief = GROOVE,textvariable = self.sugar)
        sugar_en.grid(row = 4,column = 1,ipady = 5,ipadx = 5)

        #==================Other Stuff=====================#

        F2 = LabelFrame(self.root,text = 'Others',bd = 10,relief = GROOVE,bg = bg_color,fg = "gold",font = ("times new roman",13,"bold"))
        F2.place(x = 650,y = 170,width = 325,height = 380)

        #===========Frame Content
        maza_lbl = Label(F2,font = ("times new roman",15,"bold"),fg = lbl_color,bg = bg_color,text = "Gatorade")
        maza_lbl.grid(row = 0,column = 0,padx = 10,pady = 20)
        maza_en = Entry(F2,bd = 8,relief = GROOVE,textvariable = self.gatorade)
        maza_en.grid(row = 0,column = 1,ipady = 5,ipadx = 5)

        #=======
        cock_lbl = Label(F2,font = ("times new roman",15,"bold"),fg = lbl_color,bg = bg_color,text = "Coke")
        cock_lbl.grid(row = 1,column = 0,padx = 10,pady = 20)
        cock_en = Entry(F2,bd = 8,relief = GROOVE,textvariable = self.coke)
        cock_en.grid(row = 1,column = 1,ipady = 5,ipadx = 5)

        #=======
        frooti_lbl = Label(F2,font = ("times new roman",15,"bold"),fg = lbl_color,bg = bg_color,text = "Juice")
        frooti_lbl.grid(row = 2,column = 0,padx = 10,pady = 20)
        frooti_en = Entry(F2,bd = 8,relief = GROOVE,textvariable = self.juice)
        frooti_en.grid(row = 2,column = 1,ipady = 5,ipadx = 5)

        #========
        cold_lbl = Label(F2,font = ("times new roman",15,"bold"),fg = lbl_color,bg = bg_color,text = "Waffer")
        cold_lbl.grid(row = 3,column = 0,padx = 10,pady = 20)
        cold_en = Entry(F2,bd = 8,relief = GROOVE,textvariable = self.waffer)
        cold_en.grid(row = 3,column = 1,ipady = 5,ipadx = 5)

        #============
        bis_lbl = Label(F2,font = ("times new roman",15,"bold"),fg = lbl_color,bg = bg_color,text = "Biscuits")
        bis_lbl.grid(row = 4,column = 0,padx = 10,pady = 20)
        bis_en = Entry(F2,bd = 8,relief = GROOVE,textvariable = self.biscuits)
        bis_en.grid(row = 4,column = 1,ipady = 5,ipadx = 5)

        #===================Bill Aera================#
        F3 = Label(self.root,bd = 10,relief = GROOVE)
        F3.place(x = 960,y = 170,width = 325,height = 380)
        #===========
        bill_title = Label(F3,text = "ShowBill",font = ("Lucida",13,"bold"),bd= 7,relief = GROOVE)
        bill_title.pack(fill = X)

        #============
        scroll_y = Scrollbar(F3,orient = VERTICAL)
        self.txt = Text(F3,yscrollcommand = scroll_y.set)
        scroll_y.pack(side = RIGHT,fill = Y)
        scroll_y.config(command = self.txt.yview)
        self.txt.pack(fill = BOTH,expand = 1)

        #===========Buttons Frame=============#
        F4 = LabelFrame(self.root,text = 'Bill Menu',bd = 10,relief = GROOVE,bg = bg_color,fg = "gold",font = ("times new roman",13,"bold"))
        F4.place(x = 0,y = 555,relwidth = 1,height = 145)

        #===================
        cosm_lbl = Label(F4,font = ("times new roman",12,"bold"),fg = lbl_color,bg = bg_color,text = "Total Food")
        cosm_lbl.grid(row = 0,column = 0,padx = 10,pady = 0)
        cosm_en = Entry(F4,bd = 8,relief = GROOVE,textvariable = self.total_food)
        cosm_en.grid(row = 0,column = 1,ipady = 2,ipadx = 5)

        #===================
        gro_lbl = Label(F4,font = ("times new roman",12,"bold"),fg = lbl_color,bg = bg_color,text = "Total Grocery")
        gro_lbl.grid(row = 1,column = 0,padx = 10,pady = 3)
        gro_en = Entry(F4,bd = 8,relief = GROOVE,textvariable = self.total_grocery)
        gro_en.grid(row = 1,column = 1,ipady = 2,ipadx = 5)

        #================
        oth_lbl = Label(F4,font = ("times new roman",12,"bold"),fg = lbl_color,bg = bg_color,text = "Others Total")
        oth_lbl.grid(row = 2,column = 0,padx = 10,pady = 3)
        oth_en = Entry(F4,bd = 8,relief = GROOVE,textvariable = self.total_other)
        oth_en.grid(row = 2,column = 1,ipady = 2,ipadx = 5)

        #================
        cosmt_lbl = Label(F4,font = ("times new roman",12,"bold"),fg = lbl_color,bg = bg_color,text = "Food Tax")
        cosmt_lbl.grid(row = 0,column = 2,padx = 30,pady = 0)
        cosmt_en = Entry(F4,bd = 8,relief = GROOVE,textvariable = self.tax_cos)
        cosmt_en.grid(row = 0,column = 3,ipady = 2,ipadx = 5)

        #=================
        grot_lbl = Label(F4,font = ("times new roman",12,"bold"),fg = lbl_color,bg = bg_color,text = "Grocery Tax")
        grot_lbl.grid(row = 1,column = 2,padx = 30,pady = 3)
        grot_en = Entry(F4,bd = 8,relief = GROOVE,textvariable = self.tax_groc)
        grot_en.grid(row = 1,column = 3,ipady = 2,ipadx = 5)

        #==================
        otht_lbl = Label(F4,font = ("times new roman",12,"bold"),fg = lbl_color,bg = bg_color,text = "Others Tax")
        otht_lbl.grid(row = 2,column = 2,padx = 10,pady = 3)
        otht_en = Entry(F4,bd = 8,relief = GROOVE,textvariable = self.tax_other)
        otht_en.grid(row = 2,column = 3,ipady = 2,ipadx = 5)
        
       
        
        #====================tổng================
        total_btn = Button(F4,text = "Total",bg = bg_color,fg = fg_color,font=("lucida",12,"bold"),bd = 7,relief = GROOVE,command = self.total)
        total_btn.grid(row = 1,column = 4,ipadx = 30,padx = 15)

        #========================nút=============
        genbill_btn = Button(F4,text = "Show Bill",bg = bg_color,fg = fg_color,font=("lucida",13,"bold"),bd = 7,relief = GROOVE,command = self.bill_area)
        genbill_btn.grid(row = 1,column = 5,ipadx = 20,padx = 5)

        # Tạo nút để kích hoạt hàm gửi email
        email_button = Button(F4, text='Email', font=('Arial', 12, 'bold'), bg='green', fg='white', bd=7, relief = GROOVE,command=self.send_email)
        email_button.grid(row =0,column = 4,ipadx = 30,padx=5)  # Sử dụng pack hoặc grid theo cách bố trí
        
        #===============nút thủ tong danh thu====================
        tong_lbl = Button(F4,text = "REVENUE", font=('Arial', 12, 'bold'), bg='green', fg='white',bd=7,relief = GROOVE,command=self.show_revenue_page)
        tong_lbl.grid(row =0,column = 6,ipadx = 20,padx=5)
        #========================================
        seach_lbl = Button(F4,text = "SEACH", font=('Arial', 12, 'bold'), bg='green', fg='white',bd=7,relief = GROOVE,command=self.show_Database)
        seach_lbl.grid(row =0,column = 8,ipadx = 30,padx=5)

       


        #======================nút IN BIll RA TEXT==================
        save_bill = Button(F4,text = "IN Bill",bg = bg_color,fg = fg_color,font=("lucida",12,"bold"),bd = 7,relief = GROOVE,command = self.print_bill)
        save_bill.grid(row = 0,column = 5,ipadx = 30,padx = 5)
        #====================
        clear_btn = Button(F4,text = "Clear",bg = bg_color,fg = fg_color,font=("lucida",12,"bold"),bd = 7,relief = GROOVE,command = self.clear)
        clear_btn.grid(row = 1,column = 6,ipadx = 40,padx = 5)
        #======================
        exit_btn = Button(F4,text = "Exit",bg = bg_color,fg = fg_color,font=("lucida",12,"bold"),bd = 7,relief = GROOVE,command = self.exit)
        exit_btn.grid(row = 1,column = 8,ipadx = 40)
    
    
    #===========================================================================================================
     
   
    def show_revenue_page(self):
        # Chức năng hiển thị trang doanh thu
        revenue_window = Toplevel(self.root)
        revenue_page = DanhThu(revenue_window)
    #=======hiện thị nút seach kết nối với databaseapp=======
    def show_Database(self):
        # Chức năng hiển thị trang doanh thu
        
        seach_window = Toplevel(self.root)
        seach_page = DatabaseApp(seach_window)
          
#============================================================================================================================================================
# Định nghĩa hàm gửi email
    # Hàm gửi email qua SMTP
    def send_email_via_smtp(self, email_to, subject, body):
        smtp_server = 'smtp.gmail.com'
        smtp_port = 587
        email_from = "zBachHoaminiz@gmail.com"  # Đổi sang địa chỉ email của bạn
        email_password = "mnzw yarr lynw ubma"  # Mật khẩu ứng dụng

        msg = MIMEMultipart()
        msg['From'] = email_from
        msg['To'] = email_to
        msg['Subject'] = subject

        msg.attach(MIMEText(body, 'plain'))

        try:
            server = smtplib.SMTP(smtp_server, smtp_port)
            server.starttls()  # Kích hoạt TLS
            server.login(email_from, email_password)

            server.send_message(msg)  # Gửi email

            server.quit()  # Đóng kết nối

            return True  # Email gửi thành công
        except Exception as e:
            print(f"Error sending email: {e}")
            return False

# Định nghĩa hàm send_email để mở một cửa sổ mới và nhập email cùng với nội dung
    def send_email(self):
        email_address = self.e_email.get()  # Lấy giá trị email từ Entry

        email_window = Toplevel(self.root)
        email_window.config(bg='green')

        recipient_frame = LabelFrame(email_window, text='RECIPIENT', font=('Arial', 16, 'bold'), bd=6, bg='green', fg='yellow')
        recipient_frame.grid(row=0, column=0, padx=10, pady=10, sticky='w') 

    # Nhãn và Text cho email của người nhận
        receiver_label = Label(recipient_frame, text="Email của người nhận", font=('Arial', 14, 'bold'), bg='green', fg='white')
        receiver_label.grid(row=0, column=0, padx=10, pady=10)

        receiver_entry = Text(recipient_frame, font=('Arial', 14, 'bold'), bd=2, relief=GROOVE, width=42, height=1)
        receiver_entry.insert(END, email_address)  # Thêm địa chỉ email
        receiver_entry.grid(row=0, column=1, padx=10, pady=10)

    # Nhãn cho tin nhắn và Text area cho tin nhắn
        message_label = Label(recipient_frame, text="Message", font=('Arial', 14, 'bold'), bg='green', fg='white')
        message_label.grid(row=1, column=0, padx=10, pady=10)

        email_textarea = Text(recipient_frame, font=('Arial', 14, 'bold'), bd=2, relief=GROOVE, width=42, height=11)
        email_textarea.grid(row=2, column=0, columnspan=2, padx=10, pady=10)

        email_textarea.insert(END, self.txt.get(1.0, END))  # Chuyển nội dung từ Text ban đầu sang đây

    # Nút gửi email
        send_button = Button(email_window, text='SEND', font=('Arial', 16, 'bold'), bg='green', fg='white', width=15, command=lambda: self.send_email_action(email_window,receiver_entry, email_textarea))
        send_button.grid(row=3, column=0, padx=10, pady=20)

    def send_email_action(self,email_window, receiver_entry, email_textarea):
        email_to = receiver_entry.get(1.0, END).strip()  # Địa chỉ email
        subject = "Hóa đơn từ BÁCH HÓA MINI"  # Chủ đề của email
        body = email_textarea.get(1.0, END)  # Nội dung của email

    # Gọi hàm gửi email
        if self.send_email_via_smtp(email_to, subject, body):
            messagebox.showinfo("Success", "Email đã gửi thành công!")  # Thông báo thành công
            email_window.destroy()
        else:
            messagebox.showerror("Error", " to send email")  # Email không được gửi

#===============================================================================================================================
    
#Function to get total prices
    def total(self):
        #=================Total Food Prices
        self.total_food_prices = (
            (self.bread.get() * 1)+
            (self.candy.get() * 3)+
            (self.hamburger.get() * 8)+
            (self.hotdog.get() * 6)+
            (self.sandwich.get() * 4)
        )
        self.total_food.set("$"+str(self.total_food_prices))
        self.tax_cos.set("$"+str(round(self.total_food_prices*0.05)))
        #====================Total Grocery Prices
        self.total_grocery_prices = (
            (self.wheat.get()*1)+
            (self.food_oil.get() * 5)+
            (self.salt.get() * 1)+
            (self.rice.get() *3)+
            (self.sugar.get() * 2)

        )
        self.total_grocery.set("$"+str(self.total_grocery_prices))
        self.tax_groc.set("$"+str(round(self.total_grocery_prices*0.05)))
        #======================Total Other Prices
        self.total_other_prices = (
            (self.gatorade.get() * 4)+
            (self.juice.get() * 2)+
            (self.coke.get() * 2)+
            (self.waffer.get() * 2)+
            (self.biscuits.get() * 2)
        )
        self.total_other.set("$"+str(self.total_other_prices))
        self.tax_other.set("$"+str(round(self.total_other_prices*0.05)))

    

#Function For Text Area
    def welcome_soft(self):
        self.txt.delete('1.0',END)
        # Lấy ngày và giờ hiện tại
        now = datetime.now()
        formatted_date_time = now.strftime("%d/%m/%Y %H:%M:%S")
        self.txt.insert(END,"       Welcome To Store's Retail\n")
    # Chèn ngày và giờ vào đầu văn bản
        self.txt.insert(END, f"\nDate and Time: {formatted_date_time}")
        self.txt.insert(END,f"\nBill No. : {str(self.c_bill_no.get())}")
        self.txt.insert(END,f"\nCustomer Name : {str(self.cus_name.get())}")
        self.txt.insert(END,f"\nPhone No. : {str(self.c_phone.get())}")
        self.txt.insert(END,f"\nEmail : {str(self.e_email.get())}")
        self.txt.insert(END,"\n===================================")
        self.txt.insert(END,"\nProduct          Qty         Price")
        self.txt.insert(END,"\n===================================")

    def print_bill(self):
    # Lấy nội dung của vùng văn bản
        bill_content = self.txt.get('1.0', END)
    
    # Mở một tệp văn bản để ghi nội dung
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, bill_content)
        pdf.output("bill.pdf")

    # Thông báo in hoặc lưu hóa đơn thành công
        messagebox.showinfo("THÔNG BÁO ", "Đã In Bill Thành Công !.")

 

#Function to clear the bill area
    def clear(self):
        self.txt.delete('1.0',END)
                # Xóa khu vực hóa đơn
        self.txt.delete('1.0', END)
    # Xóa các trường nhập
        self.cus_name.set("")
        self.c_phone.set("")
        self.c_bill_no.set("")
        self.e_email.set("")
        self.bread.set(0)
        self.candy.set(0)
        self.hamburger.set(0)
        self.hotdog.set(0)
        self.sandwich.set(0)
        self.rice.set(0)
        self.salt.set(0)
        self.food_oil.set(0)
        self.wheat.set(0)
        self.sugar.set(0)
        self.gatorade.set(0)
        self.coke.set(0)
        self.juice.set(0)
        self.waffer.set(0)
        self.biscuits.set(0)
        self.total_food.set("")
        self.total_grocery.set("")
        self.total_other.set("")
        self.tax_cos.set("")
        self.tax_groc.set("")
        self.tax_other.set("")
    # Tạo số hóa đơn mới
        new_bill_no = random.randint(1000, 9999)
        self.c_bill_no.set(str(new_bill_no))

#Add Product name , qty and price to bill area
    def bill_area(self):
        self.welcome_soft()
        if self.bread.get() != 0:
            self.txt.insert(END,f"\nBread             {self.bread.get()}           {self.bread.get() * 1}")
        if self.candy.get() != 0:
            self.txt.insert(END,f"\nCandy             {self.candy.get()}           {self.candy.get() * 3}")
        if self.hamburger.get() != 0:
            self.txt.insert(END,f"\nHamburger         {self.hamburger.get()}           {self.hamburger.get() * 8}")
        if self.hotdog.get() != 0:
            self.txt.insert(END,f"\nHotdog            {self.hotdog.get()}           {self.hotdog.get() * 6}")
        if self.sandwich.get() != 0 :
            self.txt.insert(END,f"\nSandwich          {self.sandwich.get()}           {self.sandwich.get() * 4}")
        if self.wheat.get() != 0:
            self.txt.insert(END,f"\nWheat             {self.wheat.get()}           {self.wheat.get() * 1}")
        if self.food_oil.get() != 0:
            self.txt.insert(END,f"\nFood Oil          {self.food_oil.get()}           {self.food_oil.get() * 5}")
        if self.salt.get() != 0:
            self.txt.insert(END,f"\nSalt              {self.salt.get()}           {self.salt.get() * 1}")
        if self.rice.get() != 0:
            self.txt.insert(END,f"\nRice              {self.rice.get()}           {self.rice.get() * 3}")
        if self.sugar.get() != 0:
            self.txt.insert(END,f"\nSugar             {self.sugar.get()}           {self.sugar.get() * 2}")
        if self.gatorade.get() != 0:
            self.txt.insert(END,f"\nGatorade          {self.gatorade.get()}           {self.gatorade.get() * 4}")
        if self.juice.get() != 0:
            self.txt.insert(END,f"\nJuice             {self.juice.get()}           {self.juice.get() * 2}")
        if self.coke.get() != 0:
            self.txt.insert(END,f"\nCoke              {self.coke.get()}           {self.coke.get() * 2}")
        if self.waffer.get() != 0:
            self.txt.insert(END,f"\nWaffer            {self.waffer.get()}           {self.waffer.get() * 2}")
        if self.biscuits.get() != 0:
            self.txt.insert(END,f"\nBiscuits          {self.biscuits.get()}           {self.biscuits.get() * 2}")
        self.txt.insert(END,"\n===================================")
        self.txt.insert(END,f"\n                      Total : ${self.total_food_prices+self.total_grocery_prices+self.total_other_prices+self.total_food_prices * 0.05+self.total_grocery_prices * 0.05+self.total_other_prices * 0.05}")
         
        # Tính tổng số tiền
        total_amount = self.total_food_prices + self.total_grocery_prices + self.total_other_prices + \
                   self.total_food_prices * 0.05 + self.total_grocery_prices * 0.05 + self.total_other_prices * 0.05

    #===========================================================================================================================================================================================================================
#================================================================================================================== 
    def ket_noi(self):
        try:
            conn = pyodbc.connect(
                Driver='{SQL Server}',  # Đảm bảo SQL Server đã cài đặt và được nhận diện bởi pyodbc
                SERVER='TEO-PC\\SQLEXPRESS',  # Thay thế bằng tên máy chủ của bạn
                Database='QLBH',  # Thay thế bằng tên cơ sở dữ liệu của bạn
                Trusted_Connection='yes'  # Sử dụng xác thực tích hợp
            )
            return conn
        except pyodbc.Error as e:
            messagebox.showerror("Lỗi kết nối", f"Không thể kết nối tới cơ sở dữ liệu: {e}")
            return None
    def save_to_database(self):
        # Đảm bảo các giá trị tổng đã được tính toán trước khi lưu vào cơ sở dữ liệu
        self.total()  # Tính tổng trước khi lưu vào cơ sở dữ liệu

        # Kết nối tới SQL Server
        conn = self.ket_noi()
        if conn:
            try:
                cursor = conn.cursor()

                # Câu lệnh INSERT INTO
                insert_query = """
                    INSERT INTO DANHSACH 
                    (BillNo, CustomerName, PhoneNo, Email, Date, Bread, Candy, 
                    Hamburger, Hotdog, Sandwich, Wheat, FoodOil, Salt, 
                    Rice, Sugar, Gatorade, Juice, Coke, Waffer, Biscuits, Total) 
                    VALUES (?, ?, ?, ?, CONVERT(DATETIME, ?, 103) , ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """

                # Tính tổng số tiền với các biến tổng đã tính trước
                total_amount = (
                    self.total_food_prices +
                    self.total_grocery_prices +
                    self.total_other_prices +
                    self.total_food_prices * 0.05 +
                    self.total_grocery_prices * 0.05 +
                    self.total_other_prices * 0.05
                )

                # Lấy ngày hiện tại
                current_date = datetime.now().strftime("%d/%m/%Y")

                # Chuẩn bị dữ liệu để chèn vào
                data = (
                    self.c_bill_no.get(),
                    self.cus_name.get(),
                    self.c_phone.get(),
                    self.e_email.get(),
                    current_date,  # Sử dụng ngày hiện tại
                    self.bread.get(),
                    self.candy.get(),
                    self.hamburger.get(),
                    self.hotdog.get(),
                    self.sandwich.get(),
                    self.wheat.get(),
                    self.food_oil.get(),
                    self.salt.get(),
                    self.rice.get(),
                    self.sugar.get(),
                    self.gatorade.get(),
                    self.juice.get(),
                    self.coke.get(),
                    self.waffer.get(),
                    self.biscuits.get(),
                    total_amount
                )

                cursor.execute(insert_query, data)  # Thực hiện chèn dữ liệu
                conn.commit()  # Lưu các thay đổi
                messagebox.showinfo("Thông báo", "Lưu dữ liệu thành công!")  # Thông báo thành công
            except pyodbc.Error as e:
                messagebox.showerror("Lỗi truy vấn", f"Lỗi khi chèn dữ liệu: {e}")  # Xử lý ngoại lệ
            finally:
                conn.close()  # Đảm bảo đóng kết nối

    def update_datetime(self):
        current_datetime = datetime.now().strftime("%d/%m/%Y - %H:%M:%S")  # Định dạng: DD/MM/YYYY - HH:MM:SS
        self.date_time_label.config(text=current_datetime)
        self.root.after(1000, self.update_datetime)

    def exit(self):
        self.root.destroy()


#============================================================================DATABASE==================================================================
class DatabaseApp:

    def __init__(self, root):
        self.root = root
        self.root.geometry("1300x700")
        self.root.maxsize(width = 1280,height = 700)
        self.root.minsize(width = 1280,height = 700)
        self.root.title("Quản lý dữ liệu khách hàng")
        self.root.configure(background="green")
        
        self.input_entries = {}
        
        self.create_widgets()
        self.doc_du_lieu()

    def ket_noi(self):
        try:
            conn = pyodbc.connect(
                Driver='{SQL Server}',
                SERVER='TEO-PC\\SQLEXPRESS',  # Thay thế bằng tên máy chủ của bạn
                Database='QLBH',  # Thay thế bằng tên cơ sở dữ liệu của bạn
                Trusted_Connection='yes'
            )
            return conn
        except pyodbc.Error as e:
            messagebox.showerror("Lỗi kết nối", f"Không thể kết nối tới cơ sở dữ liệu: {e}")
            return None

    def doc_du_lieu(self):
        conn = self.ket_noi()
        if conn:
            try:
                cursor = conn.cursor()
                cursor.execute("SELECT * FROM DANHSACH")
                rows = cursor.fetchall()

                # Xóa dữ liệu cũ trong Treeview
                self.tree.delete(*self.tree.get_children())

                # Thêm dữ liệu mới vào Treeview
                for row in rows:
                    formatted_row = [str(field).replace(",", ";") for field in row]  # Xử lý dấu phẩy
                    self.tree.insert('', 'end', values=formatted_row)

                conn.close()
            except pyodbc.Error as e:
                messagebox.showerror("Lỗi truy vấn", f"Lỗi khi đọc dữ liệu: {e}")
                conn.close()
        else:
            messagebox.showerror("Lỗi kết nối", "Không thể kết nối tới cơ sở dữ liệu")

    def cap_nhat_truong_nhap_lieu(self, event):
        selected_item = self.tree.focus()  # Lấy dòng được chọn
        if selected_item:
            selected_values = self.tree.item(selected_item, 'values')

            # Cập nhật các trường nhập liệu với dữ liệu từ hàng được chọn
            for (entry_key, value) in zip(self.input_entries.keys(), selected_values):
                self.input_entries[entry_key].delete(0, tk.END)
                self.input_entries[entry_key].insert(0, value)

    def them_du_lieu(self):
        # Lấy dữ liệu từ các trường nhập liệu
        data = tuple(self.input_entries[key].get().strip() for key in self.input_entries.keys())

        # Kiểm tra các trường bắt buộc
        if not all(data[0:4]):  # Kiểm tra 4 trường bắt buộc
            messagebox.showerror("Lỗi", "Vui lòng điền tất cả các trường bắt buộc!")
            return

        conn = self.ket_noi()
        if conn:
            try:
                cursor = conn.cursor()
                cursor.execute("SELECT COUNT(*) FROM DANHSACH WHERE BILLNo = ?", (data[0],))
                count = cursor.fetchone()[0]

                if count > 0:
                    messagebox.showerror("Lỗi", f"BILLNo '{data[0]}' đã tồn tại!")
                    conn.close()
                    return

                # Chèn dữ liệu vào cơ sở dữ liệu
                insert_query = """
                    INSERT INTO DANHSACH 
                    (BILLNo, CustomerName, PhoneNo, Email,Date, Bread, Candy, 
                    Hamburger, Hotdog, Sandwich, Wheat, FoodOil, Salt, 
                    Rice, Sugar, Gatorade, Juice, Coke, Waffer, Biscuits, Total) 
                    VALUES (?, ?, ?,?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """
                cursor.execute(insert_query, data)  # Chèn dữ liệu
                conn.commit()  # Xác nhận thay đổi
                messagebox.showinfo("Thành công", "Dữ liệu đã được thêm thành công!")
                self.doc_du_lieu()  # Cập nhật lại Treeview
            except pyodbc.Error as e:
                messagebox.showerror("Lỗi", f"Không thể thêm dữ liệu: {e}")
            finally:
                conn.close()  # Đóng kết nối
        else:
            messagebox.showerror("Lỗi kết nối", "Không thể kết nối tới cơ sở dữ liệu")

    def sua_du_lieu(self):
        selected_item = self.tree.focus()  # Lấy dòng được chọn
        if not selected_item:
            messagebox.showerror("Lỗi", "Vui lòng chọn dòng cần sửa")
            return

        # Lấy BILLNo của dòng được chọn
        bill_no = self.tree.item(selected_item, 'values')[0]
        if not bill_no:
            messagebox.showerror("Lỗi", "Không thể xác định BILLNo từ dòng được chọn")
            return

        # Kiểm tra dữ liệu nhập liệu hợp lệ cho tất cả các trường
        try:
            # Kiểm tra dữ liệu nhập liệu của từng trường và chuyển đổi sang kiểu số nguyên
            bread = int(self.input_entries["bread_entry"].get())
            candy = int(self.input_entries["candy_entry"].get())
            hamburger = int(self.input_entries["hamburger_entry"].get())
            hotdog = int(self.input_entries["hotdog_entry"].get())
            sandwich = int(self.input_entries["sandwich_entry"].get())
            wheat = int(self.input_entries["wheat_entry"].get())
            food_oil = int(self.input_entries["food_oil_entry"].get())
            salt = int(self.input_entries["salt_entry"].get())
            rice = int(self.input_entries["rice_entry"].get())
            sugar = int(self.input_entries["sugar_entry"].get())
            gatorade = int(self.input_entries["gatorade_entry"].get())
            juice = int(self.input_entries["juice_entry"].get())
            coke = int(self.input_entries["coke_entry"].get())
            waffer = int(self.input_entries["waffer_entry"].get())
            biscuits = int(self.input_entries["biscuits_entry"].get())
        except ValueError:
            messagebox.showerror("Lỗi", "Dữ liệu nhập vào phải là số nguyên")
            return

            # Tính tổng giá của các mặt hàng
        total_food_prices = (bread * 1) + (candy * 3) + (hamburger * 8) + (hotdog * 6) + (sandwich * 4)
        total_grocery_prices = (wheat * 1) + (food_oil * 5) + (salt * 1) + (rice * 3) + (sugar * 2)
        total_other_prices = (gatorade * 4) + (juice * 2) + (coke * 2) + (waffer * 2) + (biscuits * 2)

        # Tính tổng số tiền và thuế
        total_amount = total_food_prices + total_grocery_prices + total_other_prices + \
        total_food_prices * 0.05 + total_grocery_prices * 0.05 + total_other_prices * 0.05



        # Cập nhật dữ liệu của các trường nhập liệu và tổng vào Treeview
        self.input_entries["total_entry"].delete(0, tk.END)
        self.input_entries["total_entry"].insert(0, total_amount)  # Cập nhật trường Total

        # Cập nhật dữ liệu của các trường vào cơ sở dữ liệu
        update_query = """
            UPDATE DANHSACH 
            SET CustomerName = ?, PhoneNo = ?, Email = ?, Date = ?, Bread = ?, 
                Candy = ?, Hamburger = ?, Hotdog = ?, Sandwich = ?, 
                Wheat = ?, FoodOil = ?, Salt = ?, Rice = ?, Sugar = ?, 
                Gatorade = ?, Juice = ?, Coke = ?, Waffer = ?, 
                Biscuits= ?, Total= ?
            WHERE BILLNo = ?
        """
        # Tạo tuple dữ liệu cần cập nhật
        updated_data = (
            self.input_entries["customer_name_entry"].get(),
            self.input_entries["phone_no_entry"].get(),
            self.input_entries["email_entry"].get(),
            self.input_entries["date_entry"].get(),
            bread, candy, hamburger, hotdog, sandwich,
            wheat, food_oil, salt, rice, sugar,
            gatorade, juice, coke, waffer, biscuits,
            total_amount, bill_no
        )

        # Kết nối và thực hiện cập nhật
        conn = self.ket_noi()
        if conn:
            try:
                cursor = conn.cursor()  # Tạo con trỏ SQL
                cursor.execute(update_query, updated_data)  # Thực thi câu lệnh UPDATE
                conn.commit()  # Xác nhận thay đổi
                messagebox.showinfo("Thành công", "Dữ liệu đã được cập nhật thành công!")
                self.doc_du_lieu()  # Cập nhật lại Treeview để hiển thị thay đổi
            except pyodbc.Error as e:
                # Xử lý ngoại lệ nếu có lỗi trong quá trình cập nhật
                messagebox.showerror("Lỗi cập nhật", f"Không thể cập nhật dữ liệu: {e}")
            finally:
                # Đảm bảo đóng kết nối cơ sở dữ liệu
                conn.close()
        else:
            messagebox.showerror("Lỗi kết nối", "Không thể kết nối tới cơ sở dữ liệu")

    def xoa_du_lieu(self):
        selected_item = self.tree.focus()  # Lấy dòng được chọn
        if not selected_item:
            messagebox.showerror("Lỗi", "Vui lòng chọn dòng cần xóa")
            return
        
        bill_no = self.tree.item(selected_item, 'values')[0]  # Lấy BILLNo để xác định hàng

        conn = self.ket_noi()  # Kết nối cơ sở dữ liệu
        if conn:
            try:
                delete_query = "DELETE FROM DANHSACH WHERE BILLNo = ?"
                cursor = conn.cursor()  # Tạo con trỏ SQL
                cursor.execute(delete_query, (bill_no,))  # Xóa dòng được chọn
                conn.commit()  # Xác nhận thay đổi
                messagebox.showinfo("Thành công", "Dữ liệu đã được xóa thành công!")
                self.doc_du_lieu()  # Cập nhật lại Treeview
            except pyodbc.Error as e:
                messagebox.showerror("Lỗi", f"Không thể xóa dữ liệu: {e}")
            finally:
                conn.close()
        else:
            messagebox.showerror("Lỗi kết nối", "Không thể kết nối tới cơ sở dữ liệu")

    def tim_kiem_du_lieu(self):
        search_criteria = {}

        # Lấy dữ liệu từ các trường nhập liệu liên quan đến tìm kiếm
        for field_name, entry in [
            ("BILLNo", self.input_entries["bill_no_entry"]),
            ("CustomerName", self.input_entries["customer_name_entry"]),
            ("PhoneNo", self.input_entries["phone_no_entry"]),
            ("Email", self.input_entries["email_entry"]),
            ("Date", self.input_entries["date_entry"]),
            ("Bread", self.input_entries["bread_entry"]),
            ("Candy", self.input_entries["candy_entry"]),
            ("Hamburger", self.input_entries["hamburger_entry"]),
            ("Hotdog", self.input_entries["hotdog_entry"]),
            ("Sandwich", self.input_entries["sandwich_entry"]),
            ("Wheat", self.input_entries["wheat_entry"]),
            ("FoodOil", self.input_entries["food_oil_entry"]),
            ("Salt", self.input_entries["salt_entry"]),
            ("Rice", self.input_entries["rice_entry"]),
            ("Sugar", self.input_entries["sugar_entry"]),
            ("Gatorade", self.input_entries["gatorade_entry"]),
            ("Juice", self.input_entries["juice_entry"]),
            ("Coke", self.input_entries["coke_entry"]),
            ("Waffer", self.input_entries["waffer_entry"]),
            ("Biscuits", self.input_entries["biscuits_entry"]),
            ("Total", self.input_entries["total_entry"]),
        ]:
            value = entry.get().strip()  # Lấy giá trị từ trường nhập liệu
            if value:  # Nếu có giá trị, thêm vào tiêu chí tìm kiếm
                search_criteria[field_name] = value

        # Nếu không có tiêu chí nào, hiển thị thông báo lỗi
        if not search_criteria:
            messagebox.showerror("Lỗi", "Vui lòng nhập ít nhất một tiêu chí tìm kiếm!")
            return

        # Tạo truy vấn SQL với điều kiện LIKE
        query = "SELECT * FROM DANHSACH WHERE " + " AND ".join([f"{key} LIKE ?" for key in search_criteria.keys()])

        # Tạo các tham số tìm kiếm sử dụng LIKE
        search_params = tuple(f"%{v}%" for v in search_criteria.values())

        conn = self.ket_noi()  # Kết nối cơ sở dữ liệu
        if conn:
            try:
                cursor = conn.cursor()  # Tạo con trỏ SQL
                cursor.execute(query, search_params)  # Thực thi truy vấn
                rows = cursor.fetchall()  # Lấy kết quả

                # Xóa dữ liệu cũ trong Treeview
                self.tree.delete(*self.tree.get_children())

                # Thêm dữ liệu mới vào Treeview
                for row in rows:
                    formatted_row = [str(field).replace(",", ";") for field in row]  # Định dạng giá trị
                    self.tree.insert('', 'end', values=formatted_row)

                conn.close()  # Đóng kết nối
            except pyodbc.Error as e:
                messagebox.showerror("Lỗi truy vấn", f"Lỗi khi tìm kiếm dữ liệu: {e}")
                conn.close()
        else:
            messagebox.showerror("Lỗi kết nối", "Không thể kết nối tới cơ sở dữ liệu")

    def create_widgets(self):
        # Khung chứa các trường nhập liệu
        input_frame = tk.LabelFrame(
            self.root,
            text="Thông tin khách hàng",
            font=("Times New Roman", 12, "bold"),
            fg="gold",
            bg="green",
            relief=GROOVE,
            bd=10,
        )
        input_frame.pack(fill=tk.X, padx=10, pady=10)

        input_fields = [
            ("BILLNo", "bill_no_entry"),
            ("CustomerName", "customer_name_entry"),
            ("PhoneNo", "phone_no_entry"),
            ("Email", "email_entry"),
            ("Date", "date_entry"),
            ("Bread", "bread_entry"),
            ("Candy", "candy_entry"),
            ("Hamburger", "hamburger_entry"),
            ("Hotdog", "hotdog_entry"),
            ("Sandwich", "sandwich_entry"),
            ("Wheat", "wheat_entry"),
            ("FoodOil", "food_oil_entry"),
            ("Salt", "salt_entry"),
            ("Rice", "rice_entry"),
            ("Sugar", "sugar_entry"),
            ("Gatorade", "gatorade_entry"),
            ("Juice", "juice_entry"),
            ("Coke", "coke_entry"),
            ("Waffer", "waffer_entry"),
            ("Biscuits", "biscuits_entry"),
            ("Total", "total_entry"),
        ]

        # Tạo các trường nhập liệu
        for i, (label_text, entry_key) in enumerate(input_fields):
            row = i // 3  # Mỗi hàng có 3 cột
            column = i % 3  # Chỉ số cột
            
            tk.Label(input_frame, text=label_text, bg="green", fg="white").grid(row=row, column=column * 2, padx=5, pady=5)
            entry_var = ttk.Entry(input_frame)
            entry_var.grid(row=row, column=column * 2 + 1, padx=5, pady=5)
            
            self.input_entries[entry_key] = entry_var  # Lưu các trường nhập liệu

        # Khung chứa các nút điều khiển
        button_frame = tk.Frame(self.root, bg="green")
        button_frame.pack(fill=tk.X, padx=10, pady=10)

        # Các nút điều khiển
        add_button = tk.Button(
            button_frame,relief = GROOVE,bg = 'green',fg = 'white',font=("lucida",12,"bold"),bd = 7,
            text="Thêm",
            command=self.them_du_lieu,
        )
        edit_button = tk.Button(
            button_frame,relief = GROOVE,bg = 'green',fg = 'white',font=("lucida",12,"bold"),bd = 7,
            text="Sửa",
            command=self.sua_du_lieu,
        )
        delete_button = tk.Button(
            button_frame,relief = GROOVE,bg = 'green',fg = 'white',font=("lucida",12,"bold"),bd = 7,
            text="Xóa",
            command=self.xoa_du_lieu,
        )
        search_button = tk.Button(
            button_frame,
            text="Tìm kiếm",relief = GROOVE,bg = 'green',fg = 'white',font=("lucida",12,"bold"),bd = 7,
            command=self.tim_kiem_du_lieu,
        )
        reset_button = tk.Button(
            button_frame,relief = GROOVE,bg = 'green',fg = 'white',font=("lucida",12,"bold"),bd = 7,
            text="Reset",
            command=lambda: (
                [entry.delete(0, tk.END) for entry in self.input_entries.values()],  # Xóa các trường nhập liệu
                self.doc_du_lieu(),  # Đọc lại dữ liệu
            ),
        )
        
        baocao_button = tk.Button(
            button_frame,relief = GROOVE,bg = 'green',fg = 'white',font=("lucida",12,"bold"),bd = 7,
            text="Baocao",
            command=self.create_report,
        )

        exit_button = tk.Button(
            button_frame,relief = GROOVE,bg = 'green',fg = 'white',font=("lucida",12,"bold"),bd = 7,
            text="Exit",
            command=self.exit,
        )
        

        # Đảm bảo các nút được sắp xếp đúng cách
        add_button.pack(side=tk.LEFT, padx=10, pady=5)
        edit_button.pack(side=tk.LEFT, padx=10, pady=5)
        delete_button.pack(side=tk.LEFT, padx=10, pady=5)
        search_button.pack(side=tk.LEFT, padx=10, pady=5)
        reset_button.pack(side=tk.LEFT, padx=10, pady=5)
        baocao_button.pack(side=tk.LEFT, padx=10, pady=5)
        exit_button.pack(side=tk.LEFT, padx=10, pady=5)
        # Khung chứa Treeview và thanh cuộn
        tree_frame = tk.LabelFrame(
            self.root,
            text="Bảng dữ liệu",
            font=("Times New Roman", 12, "bold"),
            fg="gold",
            bg="green",
            relief=tk.GROOVE,
            bd=10,
        )
        tree_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Tạo thanh cuộn ngang
        horizontal_scrollbar = tk.Scrollbar(tree_frame, orient=tk.HORIZONTAL)
        horizontal_scrollbar.pack(side=tk.BOTTOM, fill=tk.X)  # Thanh cuộn ngang ở dưới

        # Tạo `Treeview` với liên kết đến thanh cuộn ngang
        self.tree = ttk.Treeview(
            tree_frame,
            columns=[
                "BILLNo",
                "CustomerName",
                "PhoneNo",
                "Email",
                "Date",
                "Bread",
                "Candy",
                "Hamburger",
                "Hotdog",
                "Sandwich",
                "Wheat",
                "FoodOil",
                "Salt",
                "Rice",
                "Sugar",
                "Gatorade",
                "Juice",
                "Coke",
                "Waffer",
                "Biscuits",
                "Total",
            ],
            show="headings",
            xscrollcommand=horizontal_scrollbar.set,  # Liên kết thanh cuộn ngang
        )

        # Liên kết thanh cuộn với `Treeview`
        horizontal_scrollbar.config(command=self.tree.xview)  # Đặt thanh cuộn ngang cho `Treeview`

        # Đặt tiêu đề cột và kích thước cột
        for col in self.tree["columns"]:
            self.tree.heading(col, text=col)  # Tiêu đề cột
            self.tree.column(col, width=60)  # Đặt chiều rộng cột  

        self.tree.pack(expand=True, fill=tk.BOTH)  # Đặt Treeview trong bố cục

        # Thêm sự kiện lắng nghe khi một hàng trong Treeview được chọn
        self.tree.bind("<<TreeviewSelect>>", self.cap_nhat_truong_nhap_lieu)

#------------------------------------------------------------------------------------------------------------------------
    def fetch_data(self):  # Thêm `self` vào đây
        try:
            conn = pyodbc.connect(
                Driver='{SQL Server}',
                SERVER='TEO-PC\\SQLEXPRESS',  # Thay thế bằng tên máy chủ của bạn
                Database='QLBH',  # Thay thế bằng tên cơ sở dữ liệu của bạn
                Trusted_Connection='yes'
            )
        
            cursor = conn.cursor()
            cursor.execute("SELECT Bread, Candy, Hamburger, Hotdog, Sandwich, Wheat, FoodOil, Salt, Rice, Sugar, Gatorade, Juice, Coke, Waffer, Biscuits, Date, Total FROM DANHSACH")
            rows = cursor.fetchall()
            conn.close()
            return rows

        except pyodbc.Error as e:
            print("Lỗi khi kết nối cơ sở dữ liệu:", e)
            return None

    def process_data(self,rows):
        items = ['Bread', 'Candy', 'Hamburger', 'Hotdog', 'Sandwich', 'Wheat', 'FoodOil', 'Salt', 'Rice', 'Sugar', 'Gatorade', 'Juice', 'Coke', 'Waffer', 'Biscuits']
    
    # Initialize dictionaries
        item_sales = defaultdict(int)
        monthly_sales = defaultdict(float)
        total_annual_revenue = 0
    
        for row in rows:
        # Item sales
            for i, item in enumerate(items):
                item_sales[item] += row[i]
        
        # Monthly sales and total annual revenue
            date, total = row[-2], row[-1]
            month = datetime.strptime(date, '%Y-%m-%d').strftime('%m')
            monthly_sales[month] += float(total)
            total_annual_revenue += float(total)
    
    # Find max and min item sales
        max_item = max(item_sales, key=item_sales.get)
        min_item = min(item_sales, key=item_sales.get)
    
    # Find max and min monthly sales
        max_month = max(monthly_sales, key=monthly_sales.get)
        min_month = min(monthly_sales, key=monthly_sales.get)
    
        data = {
            'max_item': max_item,
            'min_item': min_item,
            'max_month': max_month,
            'min_month': min_month,
            'total_annual_revenue': total_annual_revenue,
            'item_sales': item_sales,
            'monthly_sales': monthly_sales
        }
        return data

    def create_report(self):
        rows = self.fetch_data()
        if rows:
            data = self.process_data(rows)
        else:
            messagebox.showerror("Lỗi", "Không thể lấy dữ liệu từ cơ sở dữ liệu.")
        document = Document()
        document.add_heading('BÁO CÁO DOANH THU', 0)
    # Thêm phần tổng quan
        document.add_heading('I. Tổng Quan', level=1)
        document.add_paragraph('Tên Công Ty: Bách Hóa Mini')
        document.add_paragraph('Địa Chỉ: TP.Hồ Chí Minh')
        document.add_paragraph('Thời Kỳ Báo Cáo: 2024')

        document.add_heading('II. Thống Kê Doanh Thu Số Lượng Hàng Bán', level=1)
    # Section 1: Sản phẩm được mua nhiều nhất và ít nhất
        document.add_heading('1. Sản phẩm được mua nhiều nhất và ít nhất', level=2)
        document.add_paragraph(f"Sản phẩm được mua nhiều nhất: {data['max_item']} (Số lượng: {data['item_sales'][data['max_item']]})", style='List Bullet')
        document.add_paragraph(f"Sản phẩm được mua ít nhất: {data['min_item']} (Số lượng: {data['item_sales'][data['min_item']]})", style='List Bullet')

    # Section 2: Tổng doanh thu tháng bán nhiều nhất và ít nhất
        document.add_heading('2. Tổng doanh thu tháng bán nhiều nhất và ít nhất', level=2)
        document.add_paragraph(f"Tháng bán nhiều nhất: {data['max_month']} (Tổng doanh thu: {data['monthly_sales'][data['max_month']]:,.2f} $)", style='List Bullet')
        document.add_paragraph(f"Tháng bán ít nhất: {data['min_month']} (Tổng doanh thu: {data['monthly_sales'][data['min_month']]:,.2f} $)", style='List Bullet')

    # Section 3: Tổng doanh thu trên 1 năm
        document.add_heading('3. Tổng doanh thu trên 1 năm', level=2)
        document.add_paragraph(f"Tổng doanh thu trong năm: {data['total_annual_revenue']:,.2f} $", style='List Bullet')
    # Thêm phần kết luận
        document.add_heading('IV. Kết Luận', level=1)
        document.add_paragraph('Trên đây là báo cáo tổng kết doanh thu số lượng hàng bán được nhiều nhất và ít nhất cùng tổng doanh thu trong năm của Bách Hóa Mini. Những thông tin này cung cấp cái nhìn tổng quan về hiệu suất kinh doanh của công ty trong năm qua và có thể được sử dụng để điều chỉnh chiến lược kinh doanh trong tương lai.')
        document.add_paragraph('Xin cảm ơn!')

    
    # Thêm chữ ký bằng cách sử dụng bảng để căn giữa phải
        table = document.add_table(rows=10, cols=1)
        table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        cell.width = Inches(2)

    # Thêm khung cho chữ ký và hình ảnh
        signature_paragraph = cell.paragraphs[0]
        signature_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        signature_paragraph.paragraph_format.left_indent = Inches(1.5)  # Điều chỉnh vị trí chữ ký sang phải

        signature_run = signature_paragraph.add_run("Kí tên\n\n")
        signature_run.font.size = Pt(12)  # Chỉnh kích thước chữ ký

    # Thêm hình ảnh ký tên và điều chỉnh kích thước
        picture_run = signature_paragraph.add_run()
        picture_run.add_picture('kyten.png', width=Inches(1.5))  # Điều chỉnh kích thước hình ảnh tại đây

    # Đảm bảo font chữ của đoạn văn có kích thước phù hợp
        for run in signature_paragraph.runs:
            font = run.font
            font.size = Pt(12)   
    # Save the document
        document.save('BaoCaoDoanhThu.docx')
        messagebox.showinfo("Báo cáo Thành Công", "Báo cáo đã được tạo thành công!")
    def exit(self):
        self.root.destroy()

if __name__ == "__main__":
    root = Tk()
    app = LoadingScreen(root)
    root.mainloop()